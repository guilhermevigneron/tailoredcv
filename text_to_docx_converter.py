from docx import Document
from docx.shared import Pt, Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re
import sys
import subprocess

def create_cv_styles(document: Document) -> None:
    """
    Creates and configures custom styles for the CV document.
    
    Custom styles include:
      - Normal: Default body text (Arial, 11pt)
      - Heading 1 CV: Section headers (Arial, 14pt, bold)
      - Bullet Points CV: For list items (Arial, 11pt with indentation)
      - Value Proposition CV: Emphasis for the top section (Arial, 12pt, bold)
    """
    styles = document.styles

    # --- Normal Style (Body Text) ---
    style_normal = styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Arial'  # You can change to 'Calibri', 'Garamond', etc.
    font_normal.size = Pt(11)   # 10-12pt for body text

    # --- Heading 1 Style (Section Headers) ---
    style_heading1 = styles.add_style('Heading 1 CV', WD_STYLE_TYPE.PARAGRAPH)
    style_heading1.base_style = styles['Heading 1']  # Inherit from built-in Heading 1
    font_heading1 = style_heading1.font
    font_heading1.name = 'Arial'
    font_heading1.size = Pt(14)  # 12-14pt for headers
    font_heading1.bold = True
    paragraph_format_h1 = style_heading1.paragraph_format
    paragraph_format_h1.space_before = Pt(12)  # Spacing before headers
    paragraph_format_h1.space_after = Pt(6)    # Spacing after headers

    # --- Bullet Point Style (for lists) ---
    style_bullet = styles.add_style('Bullet Points CV', WD_STYLE_TYPE.PARAGRAPH)
    style_bullet.base_style = styles['List Bullet']  # Inherit from List Bullet style
    font_bullet = style_bullet.font
    font_bullet.name = 'Arial'
    font_bullet.size = Pt(11)  # Same as body text
    paragraph_format_bullet = style_bullet.paragraph_format
    paragraph_format_bullet.left_indent = Inches(0.25)  # Indentation for bullets
    paragraph_format_bullet.space_before = Pt(3)         # Reduced spacing
    paragraph_format_bullet.space_after = Pt(3)

    # --- Value Proposition Style (Emphasis for Top Section) ---
    style_value_prop = styles.add_style('Value Proposition CV', WD_STYLE_TYPE.PARAGRAPH)
    style_value_prop.base_style = styles['Normal']
    font_value_prop = style_value_prop.font
    font_value_prop.name = 'Arial'
    font_value_prop.size = Pt(12)  # Slightly larger for emphasis
    font_value_prop.bold = True    # Bold for impact
    paragraph_format_vp = style_value_prop.paragraph_format
    paragraph_format_vp.space_after = Pt(10)  # More space after for visual separation

def text_to_docx(text_file_path: Path, docx_file_path: Path) -> bool:
    """
    Converts a plain text file to a styled .docx CV file,
    optimized for both ATS and human readability.

    Args:
        text_file_path (Path): Path to the input text file.
        docx_file_path (Path): Path to save the output .docx file.

    Returns:
        bool: True if conversion was successful, False otherwise.
    """
    try:
        # Resolve full paths
        text_file_path = text_file_path.resolve()
        docx_file_path = docx_file_path.resolve()

        # Verify that the input text file exists
        if not text_file_path.is_file():
            print(f"Error: Text file not found at '{text_file_path}'")
            return False

        # Read the entire content of the text file
        text = text_file_path.read_text(encoding='utf-8')

        # Create a new Document and adjust section margins
        document = Document()
        for section in document.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Apply custom CV styles to the document
        create_cv_styles(document)

        # Split text into paragraphs based on one or more blank lines
        paragraphs = re.split(r'\n\s*\n', text.strip())
        is_value_proposition_section = True  # Flag for first section styling

        for para_text in paragraphs:
            stripped_text = para_text.strip()
            if not stripped_text:
                continue

            # Apply the "Value Proposition" style to the first section only.
            if is_value_proposition_section:
                document.add_paragraph(stripped_text, style='Value Proposition CV')
                is_value_proposition_section = False
            # Detect section headers based on common keywords (case-insensitive)
            elif stripped_text.upper().startswith(('SUMMARY', 'PROFILE', 'EXPERIENCE', 'EDUCATION', 
                                                    'SKILLS', 'PROJECTS', 'AWARDS', 'CERTIFICATIONS')):
                p = document.add_paragraph(stripped_text, style='Heading 1 CV')
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # Ensure left alignment for headers
            # Detect bullet points (adjust symbols as needed)
            elif stripped_text.startswith(('- ', '* ', '• ')):
                document.add_paragraph(stripped_text, style='Bullet Points CV')
            else:
                document.add_paragraph(stripped_text, style='Normal')

        # Ensure the output directory exists before saving
        docx_file_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(docx_file_path))
        print(f"Successfully converted '{text_file_path}' to '{docx_file_path}'")
        return True

    except PermissionError as pe:
        if pe.errno == 13:
            print(f"Permission denied: The file '{docx_file_path}' may be open in another application. "
                  "Please close it and try again.")
        else:
            print(f"Permission error occurred: {pe}")
        return False

    except Exception as e:
        print(f"An error occurred during the conversion: {e}")
        return False

def open_file(file_path: Path) -> None:
    """
    Opens the file using the default application based on the operating system.
    """
    try:
        if sys.platform.startswith('win'):
            import os
            os.startfile(str(file_path))
        elif sys.platform.startswith('darwin'):
            subprocess.run(['open', str(file_path)])
        else:
            subprocess.run(['xdg-open', str(file_path)])
    except Exception as e:
        print(f"Error opening the file: {e}")

if __name__ == "__main__":
    try:
        script_dir = Path(__file__).parent.resolve()
    except NameError:
        script_dir = Path.cwd()

    input_text_file = script_dir / "tailored_cv_text.txt"
    output_docx_file = script_dir / "tailored_cv.docx"

    if text_to_docx(input_text_file, output_docx_file):
        if output_docx_file.is_file():
            open_file(output_docx_file)
